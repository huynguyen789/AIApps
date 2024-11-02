import streamlit as st
import os
import asyncio
import json
import re
from openai import AsyncOpenAI
from anthropic import AsyncAnthropic
import streamlit_mermaid as stmd
import base64
import requests
import google.generativeai as genai
from docx import Document
import pypandoc
import tempfile
import anthropic
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi
from requests.exceptions import Timeout
from datetime import datetime
import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
import glob
from dotenv import load_dotenv

from unstructured_client import UnstructuredClient
from unstructured_client.models import shared, operations
from unstructured.staging.base import dict_to_elements
from unstructured.chunking.title import chunk_by_title
from langchain.schema import Document
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS

from typing import List, Tuple, Optional
import re
from rank_bm25 import BM25Okapi
import tempfile
from openai import OpenAI
# At the top of your file after imports
import logging
logging.getLogger("openai").setLevel(logging.INFO)  # Change from ERROR to INFO to see the logs

from docx import Document as DocxDocument  # For Word document handling

# Initialize clients
openai_client = AsyncOpenAI(api_key=st.secrets["OPENAI_API_KEY"])
anthropic_client = AsyncAnthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])

# MODEL HANDLING
def setup_gemini(model_variant: str = "flash"):
    api_key = st.secrets["GOOGLE_API_KEY"]
    genai.configure(api_key=api_key)
    
    generation_config = {
        "temperature": 0,
        "max_output_tokens": 8192,
    }
    
    safety_settings = [
        {"category": "HARM_CATEGORY_DANGEROUS", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
    ]
    
    model_name = "gemini-1.5-flash-002" if model_variant == "flash" else "gemini-1.5-pro-002"
    
    return genai.GenerativeModel(
        model_name=model_name,
        generation_config=generation_config,
        safety_settings=safety_settings
    )

def get_model(model_name: str):
    model_map = {
        "gemini-flash": lambda: setup_gemini("flash"),
        "gemini-pro": lambda: setup_gemini("pro"),
        "claude": lambda: anthropic_client,
        "gpt4": lambda: openai_client,
    }
    
    if model_name not in model_map:
        available_models = list(model_map.keys())
        raise ValueError(f"Unsupported model: {model_name}. Available models: {available_models}")
        
    return model_map[model_name]()

async def generate_response(model_name: str, prompt: str, system_prompt: Optional[str] = None):
    model = get_model(model_name)
    
    if isinstance(model, genai.GenerativeModel):  # Gemini (both variants)
        response = model.generate_content(prompt, stream=True)
        for chunk in response:
            if chunk.text:
                yield chunk.text
                
    elif isinstance(model, AsyncAnthropic):  # Claude
        messages = [{"role": "user", "content": prompt}]
        if system_prompt:
            messages.insert(0, {"role": "system", "content": system_prompt})
            
        async with model.messages.stream(
            model="claude-3-5-sonnet-20241022",
            max_tokens=4096,
            temperature=0,
            messages=messages
        ) as stream:
            async for text in stream.text_stream:
                yield text
                
    elif isinstance(model, AsyncOpenAI):  # GPT-4
        messages = [{"role": "user", "content": prompt}]
        if system_prompt:
            messages.insert(0, {"role": "system", "content": system_prompt})
            
        stream = await model.chat.completions.create(
            model="gpt-4o",
            messages=messages,
            max_tokens=4096,
            temperature=0,
            stream=True
        )
        async for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content

# Then update your existing functions to use the new generate_response function:

async def generate_job_description(job_title, additional_requirements, is_pws):
    main_prompt = load_prompt('job_description.txt')
    
    pws_instruction = "This is a PWS workflow. Follow the languages/wordings in the requirements strictly!!." if is_pws else ""
    formatted_prompt = f"Job title: {job_title}. {additional_requirements} {pws_instruction} \n\n {main_prompt}"
    
    async for content in generate_response("gemini-flash", formatted_prompt):
        formatted_text = content.replace('\n', '  \n')
        yield formatted_text

async def improve_job_description(original_jd, feedback, job_title, additional_requirements):
    improve_prompt = load_prompt('improve_job_description.txt')
    improve_input = f"""
Original Job Title: {job_title}
Additional Requirements: {additional_requirements}

Original Job Description:
{original_jd}

User Feedback:
{feedback}
"""
    async for content in generate_response("gpt4", improve_input, system_prompt=improve_prompt):
        yield content

def load_prompt(filename):
    with open(os.path.join('./prompts', filename), 'r') as file:
        return file.read()
#########################################################

#Prompt generation
async def get_model_answer(instruction, prompt, model_name):
    if "claude" in model_name.lower():
        message = await anthropic_client.messages.create(
            model=model_name,
            max_tokens=3000,
            temperature=0,
            system=prompt,
            messages=[
                {
                    "role": "user",
                    "content": instruction
                }
            ]
        )
        return message.content[0].text
    else:  # OpenAI models
        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": instruction}
        ]
        response = await openai_client.chat.completions.create(
            model=model_name,
            messages=messages,
            temperature=0
        )
        return response.choices[0].message.content

def remove_analysis(prompt):
    return re.sub(r'<analysis>.*?</analysis>', '', prompt, flags=re.DOTALL)

async def prompt_generator(user_request):
    optimizer_model = "claude-3-5-sonnet-20241022"  # or "gpt-4-0125-preview"
    
    if 'current_prompt' not in st.session_state:
        st.session_state.current_prompt = None
    if 'user_feedback' not in st.session_state:
        st.session_state.user_feedback = ""
    
    # Generate initial prompt if it doesn't exist
    if st.session_state.current_prompt is None:
        initial_improvement_prompt = load_prompt('create_adv_prompt.txt')
        initial_improvement_prompt = initial_improvement_prompt.format(
            user_request=user_request,
            user_feedback="No user feedback yet",
            current_prompt="No current prompt yet."
        )
        
        with st.spinner("Generating initial prompt..."):
            initial_prompt = await get_model_answer(initial_improvement_prompt, "", optimizer_model)
        
        st.session_state.current_prompt = remove_analysis(initial_prompt)

    st.subheader("Current Prompt:")
    st.text_area("Current Prompt", value=st.session_state.current_prompt, height=300, key="current_prompt_display")
    
    # Capture user feedback
    user_feedback = st.text_area("Provide feedback for improving the prompt:", value=st.session_state.user_feedback, key="user_feedback_input")
    
    col1, col2 = st.columns(2)
    improve_button = col1.button("Improve Prompt")
    reset_button = col2.button("Reset")
    
    if improve_button:
        st.session_state.user_feedback = user_feedback
        improvement_prompt = load_prompt('create_adv_prompt.txt')
        improvement_prompt = improvement_prompt.format(
            user_request=user_request,
            user_feedback=f"User feedback: {st.session_state.user_feedback}",
            current_prompt=st.session_state.current_prompt
        )
        
        with st.spinner("Generating improved prompt..."):
            improved_prompt = await get_model_answer(improvement_prompt, "", optimizer_model)
            st.session_state.current_prompt = remove_analysis(improved_prompt)
        
        st.session_state.user_feedback = ""
        st.rerun()
    
    if reset_button:
        st.session_state.current_prompt = None
        st.session_state.user_feedback = ""
        st.rerun()
    
    return st.session_state.current_prompt
#########################################################



async def generate_job_description(job_title, additional_requirements, is_pws):
    main_prompt = load_prompt('job_description.txt')
    
    pws_instruction = "This is a PWS workflow. Follow the languages/wordings in the requirements strictly!!." if is_pws else ""
    formatted_prompt = f"Job title: {job_title}. {additional_requirements} {pws_instruction} \n\n {main_prompt}"
    
    async for content in generate_response("gemini-flash", formatted_prompt):
        formatted_text = content.replace('\n', '  \n')
        yield formatted_text

async def improve_job_description(original_jd, feedback, job_title, additional_requirements):
    improve_prompt = load_prompt('improve_job_description.txt')
    improve_input = f"""
Original Job Title: {job_title}
Additional Requirements: {additional_requirements}

Original Job Description:
{original_jd}

User Feedback:
{feedback}
"""
    async for content in generate_response("gpt4", improve_input, system_prompt=improve_prompt):
        yield content
#########################################################


#Writing tools
async def process_text(user_input, task):
    prompts = {
        "grammar": "Correct the grammar in the following text...",
        "rewrite": "Rewrite the following text professionally...",
        "summarize": "Summarize the following text concisely...",
        "explain": "Explain the following text in two parts..."
    }
   
    user_input = f'"{user_input}"'
    
    async for content in generate_response("gpt4", user_input, system_prompt=prompts[task]):
        yield content

    # return processed_content
#########################################################


#BD Response Assistant
async def get_feedback(document, requirements, user_instructions):
    user_content = f"""
    <system_prompt>
    You are an expert in business development proposals for government with 30 years track record exceptional experience and of 95% winning rate proposal."
    </system_prompt>
    
    Requirements Document:\n{requirements}\n\n
    User's Draft:\n{document}\n\n
    User's Additional Instructions:\n{user_instructions}\n\n
    
    Provide detailed feedback on the user's draft generally and on the requirements provided. 
    Be critical, point out things that user dont see. The goal is a world-class winning proposal. 
    Pay special attention to the user's additional instructions and focus areas if provided.
    
    If the user has specifically requested feedback on their technical approach, provide a detailed analysis of the approach, including:
    1. Strengths of the current technical approach
    2. Areas for improvement or expansion
    3. Alignment with the requirements document
    4. Suggestions for enhancing the technical content
    5. Any potential innovative ideas that could set this proposal apart
    
    MAKE SURE YOU DOING AN EXCEPTIONAL JOB, I WILL GIVE YOU $100,000 BONUS THIS YEAR, IF NOT A CAT WILL DIE."""

    response = await openai_client.chat.completions.create(
        model="o1-preview",
        messages=[
            {"role": "user", "content": user_content}
        ]
    )
    return response.choices[0].message.content
#########################################################



#Text to diagram
async def generate_mermaid_diagram(description):
    prompt = """
    You are an expert in creating Mermaid diagrams. Based on the user's description, generate a Mermaid diagram code.
    Make sure the code is valid and follows Mermaid syntax. Return only the Mermaid code, without any additional text or explanations, tags, or code block markers.
    Try to design the code to fit nicely in horizontal monitor for user.
    

    
    Good output:
        graph TD
            A[Start] --> B[Process 1]
            B --> C[Process 2]
            C --> D[End]
            
            
    DO NOT INCLUDE THE "```mermaid"
    """
    
    user_input = f"Create a Mermaid diagram for: {description}"
    
    response = await anthropic_client.messages.create(
        model="claude-3-5-sonnet-20241022",
        max_tokens=3000,
        temperature=0,
        system=prompt,
        messages=[
            {
                "role": "user",
                "content": user_input
            }
        ]
    )
    return response.content[0].text

def mermaid_to_svg(mermaid_code):
    response = requests.post(
        'https://mermaid.ink/svg',
        json={'mermaid': mermaid_code}
    )
    if response.status_code == 200 and response.content.startswith(b'<svg'):
        return response.content
    else:
        return None
    
def generate_mermaid_chart(mermaid_code, format='png'):
    # Mermaid Live Editor API endpoint
    url = 'https://mermaid.ink/img/'

    # Encode the Mermaid code
    encoded_code = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('utf-8')

    # Construct the full URL
    full_url = f"{url}{encoded_code}"

    # Add format parameter if not PNG
    if format.lower() != 'png':
        full_url += f"?type={format}"

    # Send GET request to the API
    response = requests.get(full_url)

    # Check if the request was successful
    if response.status_code == 200:
        return response.content
    else:
        return None    
#########################################################
    

# Monthly Status Report Generator
def setup_ai_model(model_name: str):
    if model_name == "gemini":
        api_key = st.secrets["GOOGLE_API_KEY"]
        genai.configure(api_key=api_key)
        
        generation_config = {
            "temperature": 0,
            "max_output_tokens": 8192,
        }
        
        safety_settings = [
            {"category": "HARM_CATEGORY_DANGEROUS", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
            {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
        ]
        
        return genai.GenerativeModel(
            model_name="gemini-1.5-pro",
            generation_config=generation_config,
            safety_settings=safety_settings
        )
    elif model_name == "claude":
        return st.secrets["ANTHROPIC_API_KEY"]
    elif model_name == "gpt4":
        return openai_client
    else:
        raise ValueError(f"Unsupported model: {model_name}")

def create_word_document(markdown_content):
    doc = DocxDocument()
    
    # Set up styles
    styles = doc.styles
    
    title_style = styles['Title'] if 'Title' in styles else styles.add_style('Title', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    
    heading1_style = styles['Heading 1'] if 'Heading 1' in styles else styles.add_style('Heading 1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(14)
    heading1_style.font.bold = True
    
    heading2_style = styles['Heading 2'] if 'Heading 2' in styles else styles.add_style('Heading 2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(12)
    heading2_style.font.bold = True
    
    normal_style = styles['Normal']
    normal_style.font.size = Pt(11)
    
    # Process the Markdown content
    lines = markdown_content.split('\n')
    current_list = None
    for line in lines:
        if not line.strip():
            current_list = None
            if 'current_table' in locals():
                del current_table
            doc.add_paragraph()
            continue
        
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Title')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_paragraph(line[3:], style='Heading 1')
        elif line.startswith('### '):
            doc.add_paragraph(line[4:], style='Heading 2')
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='Heading 2')
            p.runs[0].italic = True
        elif line.startswith('- '):
            if current_list is None:
                current_list = doc.add_paragraph(style='List Bullet')
            current_list.add_run(line[2:])
            current_list.add_run('\n')
        elif line.startswith('|'):
            if 'current_table' not in locals():
                row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                current_table = doc.add_table(rows=1, cols=len(row_data))
                current_table.style = 'Table Grid'
                hdr_cells = current_table.rows[0].cells
                for i, val in enumerate(row_data):
                    hdr_cells[i].text = val
            else:
                row_data = [cell.strip() for cell in line.split('|') if cell.strip()]
                row_cells = current_table.add_row().cells
                for i, val in enumerate(row_data):
                    row_cells[i].text = val
        else:
            doc.add_paragraph(line)
    
    # Set consistent paragraph spacing
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.space_after = Pt(8)
    
    return doc

async def generate_monthly_status_report(model_name: str, master_content: str, example_content: str):
    prompt = load_prompt("monthly_status_report.txt")
    formatted_prompt = prompt.format(
        master_content=master_content,
        example_content=example_content
    )
    
    # Use Gemini Pro for more complex tasks
    if model_name == "gemini":
        model_name = "gemini-pro"  # Use Pro version for report generation
    
    async for content in generate_response(model_name, formatted_prompt):
        yield content

def read_file(file):
    if file.name.endswith('.docx'):
        # Use python-docx Document
        from docx import Document as DocxDocument
        doc = DocxDocument(file)
        full_text = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                full_text.append(element.text)
            elif element.tag.endswith('tbl'):
                table = []
                for row in element.findall('.//w:tr', namespaces=element.nsmap):
                    cells = [cell.text for cell in row.findall('.//w:t', namespaces=element.nsmap)]
                    table.append(' | '.join(cells))
                full_text.append('\n'.join(table))
        return '\n'.join(full_text)
    elif file.name.endswith('.txt'):
        return file.getvalue().decode('utf-8')
    else:
        raise ValueError(f"Unsupported file type: {file.name}")

def process_input_files(files):
    input_docs = {}
    for file in files:
        content = read_file(file)
        input_docs[file.name] = content
    
    # Format the master content with XML-like tags
    master_content = "\n\n\n".join([
        f"<file:{filename}>\n{content}\n</file:{filename}>"
        for filename, content in input_docs.items()
    ])
    
    return master_content

def save_markdown_to_file(markdown_content: str, file_path: str):
    with open(file_path, 'w') as md_file:
        md_file.write(markdown_content)

def create_docx_from_markdown(markdown_content):
    doc = DocxDocument()
    doc.add_paragraph(markdown_content)
    return doc

def convert_markdown_to_docx(markdown_file_path: str, output_file_path: str):
    pypandoc.convert_file(markdown_file_path, 'docx', outputfile=output_file_path)
#########################################################



#SEARCH TOOLS
def clean_content(soup):
    # Remove unnecessary elements
    for element in soup(['script', 'style', 'nav', 'footer', 'header']):
        element.decompose()

    # Try to find the main content area
    main_content = (
        soup.find('div', id='main-outlet') or  # Discourse forums
        soup.find('main') or
        soup.find('article') or
        soup.find('div', class_='content') or
        soup.find('div', id='content') or
        soup.body  # Fallback to entire body if no specific content area found
    )

    if main_content:
        # Extract text from relevant elements
        content = []
        for elem in main_content.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'blockquote', 'div', 'span', 'pre', 'img', 'table', 'a']):
            # Skip elements likely to contain metadata or navigation
            if 'class' in elem.attrs and any(c in ['crawler-post-meta', 'topic-category', 'nav', 'menu'] for c in elem['class']):
                continue
            
            # Preserve code blocks
            if elem.name == 'pre' or (elem.name == 'div' and 'class' in elem.attrs and 'code' in elem['class']):
                code_content = elem.get_text(strip=True)
                content.append(f"```\n{code_content}\n```")
            # Handle images
            elif elem.name == 'img' and elem.get('alt'):
                content.append(f"[Image: {elem['alt']}]")
            # Handle tables
            elif elem.name == 'table':
                table_content = []
                for row in elem.find_all('tr'):
                    row_content = [cell.get_text(strip=True) for cell in row.find_all(['th', 'td'])]
                    table_content.append(' | '.join(row_content))
                content.append('\n'.join(table_content))
            # Handle links
            elif elem.name == 'a':
                link_text = elem.get_text(strip=True)
                link_url = elem.get('href')
                if link_text and link_url:
                    content.append(f"[{link_text}]({link_url})")
            else:
                text = elem.get_text(strip=True)
                if text:
                    content.append(text)

        # Join the content
        cleaned_content = '\n\n'.join(content)
        # Remove extra whitespace, but not within code blocks
        cleaned_content = re.sub(r'(?<!```)\s+(?!```)', ' ', cleaned_content).strip()
        return cleaned_content
    else:
        return "No main content found."

async def search_and_summarize(query, model_choice, search_type, progress_callback=None):
    if progress_callback:
        await progress_callback("Searching for web content...", 0.1)
    
    # Serper API call
    url = "https://google.serper.dev/search"
    serper_api_key = st.secrets["SERPER_API_KEY"]
    
    payload = json.dumps({"q": query})
    headers = {
        'X-API-KEY': serper_api_key,
        'Content-Type': 'application/json'
    }
    
    try:
        response = requests.post(url, headers=headers, data=payload)
        response.raise_for_status()
        search_results = response.json()
    except requests.exceptions.RequestException as e:
        error_message = f"Error calling Serper API: {str(e)}"
        if hasattr(e, 'response') and e.response is not None:
            error_message += f"\nResponse content: {e.response.text}"
        st.error(error_message)
        return 0, 0, ""
    except json.JSONDecodeError as e:
        st.error(f"Error decoding JSON from Serper API: {str(e)}\nResponse content: {response.text}")
        return 0, 0, ""

    if 'organic' not in search_results:
        st.error(f"Unexpected response format from Serper API. Response: {search_results}")
        return 0, 0, ""

    if progress_callback:
        await progress_callback("Processing links...", 0.2)
    combined_content = "Processed Links:\n"
    successful_website_count = 0
    successful_youtube_count = 0
    total_links = 5 if search_type == "fast" else min(10, len(search_results['organic']))
    word_count = 0
    max_words = 50000

    for rank, result in enumerate(search_results['organic'][:total_links], 1):
        progress = 0.2 + (0.5 * rank / total_links)

        if word_count >= max_words:
            break

        if 'youtube.com' in result['link']:
            try:
                if progress_callback:
                    await progress_callback(f"Processing YouTube video {successful_youtube_count + 1} (Link {rank}/{total_links})", progress)
                video_id_match = re.search(r'(?:v=|\/)([0-9A-Za-z_-]{11}).*', result['link'])
                if video_id_match:
                    video_id = video_id_match.group(1)
                    transcript = YouTubeTranscriptApi.get_transcript(video_id)
                    transcript_text = ' '.join([entry['text'] for entry in transcript])
                    transcript_words = transcript_text.split()
                    
                    if word_count + len(transcript_words) > max_words:
                        transcript_words = transcript_words[:max_words - word_count]
                        transcript_text = ' '.join(transcript_words)
                    
                    combined_content += f"{successful_youtube_count + 1}. [YouTube] {result['link']}\n"
                    combined_content += f"<youtube_video_{successful_youtube_count+1} rank={rank} title='{result['title']}' url='{result['link']}'>\n"
                    combined_content += transcript_text
                    combined_content += f"\n</youtube_video_{successful_youtube_count+1}>\n\n"
                    successful_youtube_count += 1
                    word_count += len(transcript_words)
            except Exception as e:
                pass
        
        else:
            try:
                if progress_callback:
                    await progress_callback(f"Processing website {successful_website_count + 1} (Link {rank}/{total_links})", progress)
                page_response = requests.get(result['link'], timeout=5)
                soup = BeautifulSoup(page_response.content, 'html.parser')
                cleaned_content = clean_content(soup)
                content_words = cleaned_content.split()
                
                if word_count + len(content_words) > max_words:
                    content_words = content_words[:max_words - word_count]
                    cleaned_content = ' '.join(content_words)
                
                combined_content += f"{successful_website_count + 1}. [Website] {result['link']}\n"
                combined_content += f"<website_{successful_website_count+1} rank={rank} url='{result['link']}'>\n"
                combined_content += cleaned_content
                combined_content += f"\n</website_{successful_website_count+1}>\n\n"
                successful_website_count += 1
                word_count += len(content_words)
            except Timeout:
                pass
            except Exception as e:
                pass

    combined_content += "\n\nProcessed Content:\n"

    if progress_callback:
        await progress_callback(f"Generating LLM response... (Used {successful_website_count} websites and {successful_youtube_count} YouTube videos)", 0.7)
    
    # Load the search prompt
    search_prompt = load_prompt('search_summarize.txt')
    
    # Format the prompt with the query and combined content
    formatted_prompt = search_prompt.format(
        query=query,
        combined_content=combined_content,
        today_date=datetime.now().strftime("%Y-%m-%d")
    )
    
    # Use the selected model for generating the response
    model = setup_model(model_choice)
    response_container = st.empty()
    full_response = ""
    async for content in generate_response(model_choice, formatted_prompt):
        full_response += content
        response_container.markdown(full_response)

    if progress_callback:
        await progress_callback("Search and summarize process completed", 1.0)

    return successful_website_count, successful_youtube_count, combined_content, full_response, word_count

async def stream_response(prompt):
    response = await openai_client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a world-class search engine."},
            {"role": "user", "content": prompt}
        ],
        stream=True,
    )

    response_container = st.empty()
    full_response = ""
    async for chunk in response:
        if chunk.choices[0].delta.content:
            content = chunk.choices[0].delta.content
            full_response += content
            response_container.markdown(full_response)
#########################################################


#RAG TOOL:
# Load environment variables
load_dotenv()
openai_api_key = st.secrets["OPENAI_API_KEY"]
client = OpenAI(api_key=openai_api_key)

# Get unstructured API key
unstructured_api_key = os.getenv("UNSTRUCTURED_API_KEY")
if not unstructured_api_key:
    raise ValueError("UNSTRUCTURED_API_KEY environment variable not found")

class EnhancedRetriever:
    def __init__(self, vectorstore, documents):
        self.vectorstore = vectorstore
        self.documents = documents
        self.bm25 = self._create_bm25_index()

    def _create_bm25_index(self):
        tokenized_docs = [self._tokenize(doc.page_content) for doc in self.documents]
        return BM25Okapi(tokenized_docs)

    def _tokenize(self, text: str) -> List[str]:
        tokens = re.findall(r'\b\d+(?:\.\d+)*(?:\s+[A-Za-z]+(?:\s+[A-Za-z]+)*)?\b|\w+', text.lower())
        return tokens

    def hybrid_search(self, query: str, k: int = 4, verbose: bool = False) -> List[Tuple[float, Document]]:
        # Get k semantic results and 1 keyword result
        vector_results = self.vectorstore.similarity_search_with_score(query, k=k)
        keyword_results = self.keyword_search(query, k=1)
        
        if verbose:
            st.write(f"🔍 Semantic search returned: {len(vector_results)} results")
            st.write(f"🔑 Keyword search returned: {len(keyword_results)} results")
        
        combined_results = {}
        query_keywords = set(query.lower().split())
        
        # Always include all semantic search results
        for doc, score in vector_results:
            combined_results[doc.page_content] = {'doc': doc, 'vector_score': score, 'keyword_score': 0, 'exact_match': False}
            doc_words = set(doc.page_content.lower().split())
            if query_keywords.issubset(doc_words):
                combined_results[doc.page_content]['exact_match'] = True
        
        # Always include keyword search result, even if it creates a new entry
        for score, doc in keyword_results:
            if doc.page_content in combined_results:
                combined_results[doc.page_content]['keyword_score'] = score
            else:
                # Always add keyword result as a new entry if not already present
                combined_results[doc.page_content] = {'doc': doc, 'vector_score': 0, 'keyword_score': score, 'exact_match': False}
            
            doc_words = set(doc.page_content.lower().split())
            if query_keywords.issubset(doc_words):
                combined_results[doc.page_content]['exact_match'] = True
        
        # After processing all results, check if any exact matches were found
        exact_matches = [content for content, scores in combined_results.items() if scores['exact_match']]
        if exact_matches:
            st.write("✨ Found exact matches for query terms!")
        else:
            st.write("📝 No exact matches found for query terms")
        
        final_results = []
        for content, scores in combined_results.items():
            normalized_vector_score = 1 / (1 + scores['vector_score'])
            normalized_keyword_score = scores['keyword_score']
            exact_match_bonus = 2 if scores['exact_match'] else 0
            combined_score = (normalized_vector_score + normalized_keyword_score + exact_match_bonus) / 3
            
            # Add source information to the document metadata
            scores['doc'].metadata['search_source'] = []
            if scores['vector_score'] > 0:
                scores['doc'].metadata['search_source'].append('semantic')
            if scores['keyword_score'] > 0:
                scores['doc'].metadata['search_source'].append('keyword')
            
            final_results.append((combined_score, scores['doc']))
        
        final_results = sorted(final_results, key=lambda x: x[0], reverse=True)
        
        if verbose:
            st.write(f"📊 Final combined results: {len(final_results)} documents")
            st.write(f"🔍 First 5 results:")
            for score, doc in final_results[:5]:
                st.write(f"Score: {score:.4f}")
                st.write(f"🔎 {' & '.join(doc.metadata['search_source']).title()} Search")
                st.write(f"Content: {doc.page_content[:200]}...")
                st.write("---")
        
        return final_results

    def keyword_search(self, query: str, k: int = 4) -> List[Tuple[float, Document]]:
        tokenized_query = self._tokenize(query)
        bm25_scores = self.bm25.get_scores(tokenized_query)
        scored_docs = [(score, self.documents[i]) for i, score in enumerate(bm25_scores)]
        return sorted(scored_docs, key=lambda x: x[0], reverse=True)[:k]

def process_pdfs_and_cache(input_folder, output_folder, strategy, cache_file_path=None):
    """Process PDFs and cache results. If cache_file_path is provided, use that instead of generating one."""
    s = UnstructuredClient(
        api_key_auth=st.secrets["UNSTRUCTURED_API_KEY"],
        server_url='https://redhorse-d652ahtg.api.unstructuredapp.io'
    )

    os.makedirs(output_folder, exist_ok=True)
    
    # If no specific cache path provided, generate one from the input folder
    if cache_file_path is None:
        folder_name = os.path.basename(os.path.normpath(input_folder))
        cache_file_path = os.path.join(output_folder, f'{folder_name}_combined_content.json')
    
    # Check if cache exists first
    if os.path.exists(cache_file_path):
        st.info("Using cached version of the processed files.")
        with open(cache_file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
            
    # Only process if cache doesn't exist
    combined_content = []
    pdf_files = glob.glob(os.path.join(input_folder, "*.pdf"))
    
    if not pdf_files:
        raise ValueError(f"No PDF files found in {input_folder}")
        
    total_files = len(pdf_files)
    for idx, filename in enumerate(pdf_files, 1):
        st.write(f"Processing file {idx}/{total_files}: {os.path.basename(filename)}")
        
        try:
            with open(filename, "rb") as file:
                partition_params = shared.PartitionParameters(
                    files=shared.Files(
                        content=file.read(),
                        file_name=os.path.basename(filename),
                    ),
                    strategy=strategy,
                )
                req = operations.PartitionRequest(
                    partition_parameters=partition_params
                )
                res = s.general.partition(request=req)
                combined_content.extend(res.elements)
                st.write(f"✅ Successfully processed {os.path.basename(filename)}")
        except Exception as e:
            st.error(f"Error processing {os.path.basename(filename)}: {str(e)}")
            continue

    # Save directly to cache file if processing was successful
    if combined_content:
        with open(cache_file_path, 'w', encoding='utf-8') as f:
            json.dump(combined_content, f)
        st.success(f"Successfully processed {total_files} files and saved to cache")
    else:
        st.error("No content was successfully processed")

    return combined_content

def process_data(combined_content):
    pdf_elements = dict_to_elements(combined_content)
    elements = chunk_by_title(pdf_elements, combine_text_under_n_chars=4000, max_characters=8000, new_after_n_chars=7000, overlap=1000)
    documents = []
    for element in elements:
        metadata = element.metadata.to_dict()
        metadata.pop("languages", None)
        metadata["source"] = metadata["filename"]
        documents.append(Document(page_content=element.text, metadata=metadata))
    return documents

def organize_documents(docs):
    organized_text = ""
    for i, doc in enumerate(docs, start=1):
        source = doc.metadata.get('source', 'unknown source')
        page_number = doc.metadata.get('page_number', 'unknown page number')
        organized_text += f"Document {i}:\nSource: {source}\nPage number: {page_number}\nContent: {doc.page_content}\n\n"
    return organized_text

def generate_answer(query: str, relevant_data: str, filename: str):
    prompt = f"""
    <retrieval data>
    {relevant_data}
    </retrieval data>
    
    <user query>
    {query}
    </user query>
    
    <instructions>
    You are a world-class RAG system. Your task is to give exceptional, useful, and truthful answers, based on the user's query and the provided relevant data.

    Guidelines:
    - Provide clear, accurate answers using only the given information. If information is insufficient, acknowledge limitations.
    - Connect related sections if they appear fragmented (check section numbers)
    - Use natural, easy to understand language.
    - Structure the answer in a super nice and easy to view with title, subtitle, and bullet points.
    - Remember you are discussing content from: {filename}

    Format your response in 3 parts:
    1. Short answer: Short and concise answer. 
    2. Detailed answer: Comprehensive answer with exact words in quotation marks (as much as possible) and examples where relevant. 
    3. Sources: Reference specific documents/pages used, links. 
    
    Answer in a nice markdown format:
    </instructions>
    """

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
        ],
        stream=True,
    )
    
    for chunk in response:
        if chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content

def extract_search_keywords(query: str, filename: str = "") -> str:
    prompt = f"""
    This is a RAG system with semantic and keyword search.
    
    <context>
    You are searching through documents from: {filename}
    </context>
    
    Your task is to extract the most relevant search keywords or phrases from this query for searching through documents.
    Focus on specific terms, section numbers, or phrases that are likely to yield the most relevant results.
    If user asked for a summary, should search for 'table of content' and other words that can get good results.
    Return your answer as a comma-separated string of keywords.
    
    Consider the document context when selecting keywords. For example:
    - If the filename suggests a technical document, prioritize technical terms
    - If it's a policy document, focus on policy-related terms
    - If it's a proposal, look for proposal-specific terminology
    
    User query: {query}
    Document context: {filename}
    """
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": query}
        ],
        temperature=0
    )
    result = response.choices[0].message.content
    print(f"Extracted keyword: {result}")
    return result.strip()

def rag_query_enhanced(user_query: str, enhanced_retriever: EnhancedRetriever, k: int = 4, filename: str = "", verbose: bool = False):
    search_keywords = extract_search_keywords(user_query, filename)
    if verbose:
        st.write("🔍 Extracted keywords:", search_keywords)
    
    retrieved_docs = enhanced_retriever.hybrid_search(search_keywords, k=k, verbose=verbose)
    
    organized_text = organize_documents([doc for _, doc in retrieved_docs])
    if verbose:
        st.write("\n📝 Final Context for LLM:")
        st.write(organized_text)
    
    answer = generate_answer(user_query, organized_text, filename)
    return answer

def get_cache_folders():
    cache_dir = "./cache"
    return [f for f in os.listdir(cache_dir) if f.endswith('.json')]

def create_cache_key(files, custom_name=None):
    """Create a unique cache key based on file contents and custom name if provided"""
    # Get the names of all files, stripped of extensions
    file_names = [os.path.splitext(file.name)[0] for file in files]
    file_names.sort()  # Sort for consistency
    
    # Create a content-based hash that will be the same for the same files
    files_info = []
    for file in files:
        content_hash = hash(file.getvalue())  # Hash of file content
        file_info = f"{file.name}_{content_hash}"
        files_info.append(file_info)
    files_info.sort()  # Sort for consistency
    content_hash = hash("_".join(files_info))
    
    # Use custom name if provided, otherwise create name from file names
    if custom_name:
        files_preview = custom_name
    else:
        if len(file_names) <= 3:
            files_preview = "_".join(file_names)
        else:
            files_preview = f"{file_names[0]}_{file_names[1]}_and_{len(file_names)-2}_more"
    
    # Create final filename: readable prefix + content hash
    return f"{files_preview}_{abs(content_hash)}.json"

def process_uploaded_files(uploaded_files, custom_name=None):
    """Process multiple uploaded files and return cache file name"""
    output_folder = "./cache"
    os.makedirs(output_folder, exist_ok=True)
    
    # Create a cache key based on the uploaded files
    cache_file_name = create_cache_key(uploaded_files, custom_name)
    cache_file_path = os.path.join(output_folder, cache_file_name)

    # Check if these exact files have already been processed
    if os.path.exists(cache_file_path):
        st.info("These files have already been processed. Using existing cache.")
        return cache_file_name

    # Process the files if not in cache
    with tempfile.TemporaryDirectory() as temp_dir:
        # Save all uploaded files to temp directory
        for uploaded_file in uploaded_files:
            temp_file_path = os.path.join(temp_dir, uploaded_file.name)
            with open(temp_file_path, "wb") as f:
                f.write(uploaded_file.getvalue())
        
        # Process all PDFs in the temporary directory
        strategy = "auto"
        process_pdfs_and_cache(temp_dir, output_folder, strategy, cache_file_path)
    
    return cache_file_name

# Add this to the existing functions
def process_query(query, retriever, k, conversation_history, filename, verbose=False):
    try:
        context = "\n".join([f"User: {q}\nAI: {a}" for q, a in conversation_history])
        full_query = f"Conversation history:\n{context}\n\nUser's new query: {query}"
        return rag_query_enhanced(full_query, retriever, k=k, filename=filename, verbose=verbose)
    except Exception as e:
        st.error(f"Error processing query: {str(e)}")
        return "I apologize, but I encountered an error processing your query. Please try again."
@st.cache_resource
def initialize_retriever_from_cache(cache_file_path):
    with open(cache_file_path, 'r', encoding='utf-8') as f:
        combined_content = json.load(f)
    documents = process_data(combined_content)
    embeddings = OpenAIEmbeddings(openai_api_key=st.secrets["OPENAI_API_KEY"])
    vectorstore = FAISS.from_documents(documents, embeddings)
    return EnhancedRetriever(vectorstore, documents)


@st.cache_resource
def initialize_retriever(folder_path):
    strategy = "fast"
    combined_content = process_pdfs_and_cache(folder_path, "./cache", strategy)
    documents = process_data(combined_content)
    embeddings = OpenAIEmbeddings()
    vectorstore = FAISS.from_documents(documents, embeddings)
    return EnhancedRetriever(vectorstore, documents)

def export_conversation(conversation_history):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    export_data = ""
    for user_msg, ai_msg in conversation_history:
        export_data += f"User: {user_msg}\n\nAssistant: {ai_msg}\n\n---\n\n"
    return export_data, timestamp

def trim_conversation_history(conversation_history, max_words=50000):
    """Trim conversation history to stay under max_words limit"""
    total_words = 0
    # Count words from newest to oldest
    for i in range(len(conversation_history) - 1, -1, -1):
        user_msg, ai_msg = conversation_history[i]
        words_in_exchange = len(user_msg.split()) + len(ai_msg.split())
        total_words += words_in_exchange
        
        if total_words > max_words:
            # Keep only the messages that fit within the limit
            st.warning(f"Conversation history exceeded {max_words:,} words. Removing older messages to maintain performance.")
            return conversation_history[i+1:]
    
    return conversation_history

#MAIN APP
async def streamlit_main():
    # Disable OpenAI request logging
    logging.getLogger("openai").setLevel(logging.INFO)

    st.set_page_config(page_title="AI Assistant Tools", page_icon="🛠️", layout="wide")

    st.title("AI Assistant Tools 🛠️")

    # Create categories in the sidebar
    st.sidebar.header("Categories")
    category = st.sidebar.radio("Select Category:", [
        "Home",
        "Business Development",
        "Everyday Use"
    ])

    # Organize tools by category
    if category == "Business Development":
        tool_choice = st.sidebar.radio("Choose a tool:", [
            "BD Response Assistant",
            "Job Description Assistant",
            "Monthly Report Assistant",
            "Document Chat Assistant"
        ])
    elif category == "Everyday Use":
        tool_choice = st.sidebar.radio("Choose a tool:", [
            "Search Assistant",
            "Writing Assistant",
            "Diagram Creation Assistant",
            "Prompt Engineering Assistant"
        ])
    else:  # Home
        tool_choice = "Home"

    if tool_choice == "Home":
        st.markdown("""
            Welcome! This app provides various AI-powered assistants for internal use. 
            Choose an assistant from the sidebar to get started.

            ## Privacy Notice
            - In this app, OpenAI will not look at conversations and deletes data after 30 days
            - Chat histories are temporary and will be cleared when you close or refresh the page

            ## Usage Guidelines
            - Not approved for use with CUI/FOUO or classified data
            - Don't send personal info: email, phone number, DOB, address
            - Don't use in hiring decisions (e.g. matching resumes to jobs)
            - Each client has their own AI policies, so talk to your manager to discuss appropriate client use cases
            - Use common sense and reach out if you have questions

            ## Important Note
            ⚠️ LLMs can make up information. You should treat responses as a starting point or draft and be sure to verify any information.

            ## Feedback & Suggestions
            We're constantly improving! Please reach out if you:
            - Found a bug 🐛
            - Have ideas for improvements ✨
            - Want to request new features 🚀
            - Have any questions or concerns ❓

            **Contact:** Huy Nguyen on company Microsoft Teams
            """
        )
    elif tool_choice == "Prompt Engineering Assistant":
        st.header("Prompt Engineering Assistant 🧠")
        st.write("Generate and refine prompts for AI models.")

        user_request = st.text_input("Enter your request for prompt generation:")
        if user_request:
            await prompt_generator(user_request)
            
    elif tool_choice == "Job Description Assistant":
        st.header("Job Description Assistant 📝")
        st.write("Enter the job title and any additional requirements to generate a high-quality job description.")

        if 'job_description' not in st.session_state:
            st.session_state.job_description = ""
        if 'job_title' not in st.session_state:
            st.session_state.job_title = ""
        if 'additional_requirements' not in st.session_state:
            st.session_state.additional_requirements = ""
        if 'is_pws' not in st.session_state:
            st.session_state.is_pws = False

        job_title = st.text_input("Enter the job title:", 
                                  value=st.session_state.job_title, 
                                  placeholder="e.g., Senior Software Engineer",
                                  key="job_title_input")

        # File upload option for additional requirements
        req_file = st.file_uploader("Upload additional requirements (Word or Text file) - Optional", type=['docx', 'txt'])
        if req_file:
            st.session_state.additional_requirements = read_file(req_file)
        
        additional_requirements = st.text_area("Enter any additional requirements (optional):", 
                                               value=st.session_state.additional_requirements,
                                               placeholder="TS clearance, 5+ years of experience in Python, knowledge of machine learning, etc.",
                                               key="job_description_requirements")

        # Add PWS checkbox
        is_pws = st.checkbox("This is a PWS (model will follow PWS language strictly)", value=st.session_state.is_pws, key="is_pws_checkbox")
        st.session_state.is_pws = is_pws

        if st.button("Generate Job Description"):
            if job_title:
                with st.spinner("Generating job description..."):
                    job_description_placeholder = st.empty()
                    full_content = ""
                    async for content in generate_job_description(job_title, additional_requirements, is_pws):
                        full_content += content
                        job_description_placeholder.markdown(full_content)
                    st.session_state.job_description = full_content
                    st.session_state.job_title = job_title
                    st.session_state.additional_requirements = additional_requirements
            else:
                st.warning("Please enter a job title.")
        
        # Display the current job description
        # if st.session_state.job_description:
        #     st.markdown("### Current Job Description")
        #     st.markdown(st.session_state.job_description)
                
        feedback = st.text_area("Provide feedback to improve the job description:", placeholder="example: Follow exactly the PWS languages, tailor for CBM+ projects, 7 years of experience instead of 5, etc...")
        
        if st.button("Improve Job Description"):
            if feedback and st.session_state.job_description:
                with st.spinner("Improving job description..."):
                    improved_jd_placeholder = st.empty()
                    improved_content = ""
                    async for content in improve_job_description(st.session_state.job_description, feedback, st.session_state.job_title, st.session_state.additional_requirements):
                        improved_content += content
                        improved_jd_placeholder.markdown(improved_content)
                    st.session_state.job_description = improved_content
            elif not st.session_state.job_description:
                st.warning("Please generate a job description first.")
            else:
                st.warning("Please provide feedback to improve the job description.")
   
    elif tool_choice == "Writing Assistant":
        st.header("Writing Assistant ✍️")
        st.write("Welcome to the Writing Assistant! Here you can improve your writing with AI-powered tools. Choose from the following options:")
        st.write("1. Professional Rewrite: Enhance the professionalism of your text.")
        st.write("2. Correct Grammar: Fix grammatical errors in your text.")
        st.write("3. Summarize: Get a concise summary of your text.")
        st.write("4. Explain: Simplify and explain your text.")

        if 'user_input' not in st.session_state:
            st.session_state.user_input = ""
        if 'task' not in st.session_state:
            st.session_state.task = "Professional Rewrite"
        if 'processed_text' not in st.session_state:
            st.session_state.processed_text = ""

        # File upload option for user input
        input_file = st.file_uploader("Upload your text (Word or Text file)", type=['docx', 'txt'])
        if input_file:
            st.session_state.user_input = read_file(input_file)
        
        user_input = st.text_area("Or enter your text here:", 
                                  value=st.session_state.user_input,
                                  height=200, 
                                  key="user_input")
        
        task = st.selectbox("Choose a task:", ["Professional Rewrite", "Correct Grammar", "Summarize", "Explain"], key="task")
        
        if st.button("Process Text"):
            if user_input:
                task_map = {
                    "Correct Grammar": "grammar",
                    "Professional Rewrite": "rewrite",
                    "Summarize": "summarize",
                    "Explain": "explain"
                }
        
                with st.spinner(f"Processing text ({task.lower()})..."):
                    processed_text_placeholder = st.empty()
                    full_content = ""
                    async for content in process_text(user_input, task_map[task]):
                        full_content += content
                        processed_text_placeholder.markdown(full_content)
                    st.session_state.processed_text = full_content

        # if st.session_state.processed_text:
            # st.markdown(f"### {task} Result:")
            # st.write(st.session_state.processed_text)

    elif tool_choice == "BD Response Assistant":
        st.header("BD Response Assistant 📄")
        
        st.write("Welcome! We're here to help you refine your BD response. You can either paste your draft and requirements or upload files.")
        
        st.warning("Note: This assistant cannot review figures and tables yet. Please ensure any critical information in figures or tables is also described in text.")

        # Initialize session state variables
        if "bd_document" not in st.session_state:
            st.session_state.bd_document = ""
        if "bd_requirements" not in st.session_state:
            st.session_state.bd_requirements = ""
        if "bd_user_instructions" not in st.session_state:
            st.session_state.bd_user_instructions = ""

        # File upload option for BD response draft
        bd_file = st.file_uploader("Upload your BD Response Draft (Word or Text file)", type=['docx', 'txt'])
        if bd_file:
            st.session_state.bd_document = read_file(bd_file)
        else:
            st.session_state.bd_document = st.text_area("Or paste your BD Response Draft here:", 
                                                        value=st.session_state.bd_document, 
                                                        height=100)  

        # File upload option for requirements document
        req_file = st.file_uploader("Upload the Requirements Document (Word or Text file)", type=['docx', 'txt'])
        if req_file:
            st.session_state.bd_requirements = read_file(req_file)
        else:
            st.session_state.bd_requirements = st.text_area("Or paste the Requirements Document here:", 
                                                            value=st.session_state.bd_requirements, 
                                                            height=100)  

        # New section for user instructions or comments
        st.session_state.bd_user_instructions = st.text_area(
            "Additional Instructions or Focus Areas (optional):",
            value=st.session_state.bd_user_instructions,
            height=100,
            help="Provide any specific areas you want feedback on, or any particular aspects of the technical approach you'd like the AI to focus on. For example: 'I'm working on an early draft and want feedback on my technical approach.'"
        )

        if st.button("Get Feedback"):
            if st.session_state.bd_document.strip() == "" or st.session_state.bd_requirements.strip() == "":
                st.warning("Please provide both the BD response draft and the requirements document.")
            else:
                with st.spinner("Generating feedback..."):
                    feedback = await get_feedback(st.session_state.bd_document, st.session_state.bd_requirements, st.session_state.bd_user_instructions)
                    st.markdown(feedback)

    elif tool_choice == "Diagram Creation Assistant":
        st.header("Diagram Creation Assistant 🎨")
        st.write("Transform your ideas into visual diagrams with ease!")
        st.write("Here are some example descriptions you can use:")
        st.code("""
        1. A flowchart of making coffee
        2. A flowchart of an MLOps pipeline
        3. A flowchart of a user registration process
        """, language="markdown")

        # Initialize session state variables
        if "diagram_description" not in st.session_state:
            st.session_state.diagram_description = ""
        if "mermaid_code" not in st.session_state:
            st.session_state.mermaid_code = ""

        # Use session state for the text area
        st.session_state.diagram_description = st.text_area(
            "Describe the diagram you want to create:", 
            value=st.session_state.diagram_description,
            placeholder="e.g., A flowchart showing the steps to plan a vacation, a flowchart of a MLOps system"
        )

        if st.button("Create Diagram"):
            if st.session_state.diagram_description:
                with st.spinner("Creating your diagram..."):
                    st.session_state.mermaid_code = await generate_mermaid_diagram(st.session_state.diagram_description)

        if st.session_state.mermaid_code:
            st.subheader("Your Generated Diagram:")
            try:
                mermaid = stmd.st_mermaid(st.session_state.mermaid_code, height=800)
            except Exception as e:
                st.error(f"Oops! There was an error creating your diagram: {str(e)}")
                st.text("Technical details (for troubleshooting):")
                st.code(st.session_state.mermaid_code, language="mermaid")
            else:
                st.subheader("Diagram Code:")
                st.code(st.session_state.mermaid_code, language="mermaid")
                st.info("You can copy this code and edit the diagram in [Mermaid Live](https://mermaid.live/).")
                
                # Generate PNG image
                png_image = generate_mermaid_chart(st.session_state.mermaid_code, format='png')
                if png_image:
                    st.download_button(
                        label="Download Diagram as PNG",
                        data=png_image,
                        file_name="mermaid_diagram.png",
                        mime="image/png",
                        key="png_download"
                    )
                else:
                    st.warning("Unable to generate a downloadable PNG. You can still use the Mermaid code above.")

        if st.button("Reset"):
            st.session_state.diagram_description = ""
            st.session_state.mermaid_code = ""
            st.rerun()
    elif tool_choice == "Monthly Report Assistant":
        st.header("Monthly Report Assistant 📊")
        st.write("""
        Welcome to the Monthly Report Assistant! This tool simplifies the process of creating 
        comprehensive monthly status reports by combining individual reports into a single, cohesive document. 
        Get started by uploading your files below!
        """)
        
        # Initialize session state variables
        if 'report_content' not in st.session_state:
            st.session_state.report_content = ""
        if 'master_content' not in st.session_state:
            st.session_state.master_content = ""
        if 'files_processed' not in st.session_state:
            st.session_state.files_processed = False

        uploaded_files = st.file_uploader("Upload input files (Word or Text)", type=['docx', 'txt'], accept_multiple_files=True)

        if uploaded_files and not st.session_state.files_processed:
            st.session_state.master_content = process_input_files(uploaded_files)
            st.session_state.files_processed = True

        if st.session_state.files_processed:
            with open("./example/example.txt", 'r') as file:
                example_content = file.read()

            model_choice = st.selectbox("Choose AI model:", ["gemini-flash", "gemini-pro", "gpt4", "claude"])

            if st.button("Generate Report") or st.session_state.report_content:
                if not st.session_state.report_content:
                    report_placeholder = st.empty()
                    report_content = ""
                    
                    async for chunk in generate_monthly_status_report(model_choice, st.session_state.master_content, example_content):
                        if chunk == report_content:  # This is the full response
                            st.session_state.report_content = chunk
                        else:
                            report_content += chunk
                            report_placeholder.markdown(report_content + "▌")

            # Create Word document
            doc = create_word_document(st.session_state.report_content)
            
            # Save DOCX to BytesIO object
            docx_bio = io.BytesIO()
            doc.save(docx_bio)
            docx_bio.seek(0)

            # Provide download button for Word Report
            st.download_button(
                label="Download Word Report",
                data=docx_bio.getvalue(),
                file_name="monthly_status_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
                
            if st.button("Reset"):
                st.session_state.report_content = ""
                st.session_state.master_content = ""
                st.session_state.files_processed = False
                st.rerun()
                
            #clear the streaming output before display the report
            st.empty()
            with st.expander("### Generated Report"):
                st.markdown(st.session_state.report_content)

    elif tool_choice == "Search Assistant":
        st.header("AI-Powered Search Assistant 🔍")
        st.write("Enter a search query, and the AI will search multiple websites and YouTube videos, then provide a concise and detailed answer.")

        st.markdown("""
        ### Example Queries:
        1. "What are the latest cybersecurity requirements for DoD contractors?"
        2. "How to use gpt4o openapi api python"
        3. "Recent advancements in AI for military applications"
        """)

        # Initialize session state variables if they don't exist
        if 'search_query' not in st.session_state:
            st.session_state.search_query = ""
        if 'search_model_choice' not in st.session_state:
            st.session_state.search_model_choice = "gemini"
        if 'search_type' not in st.session_state:
            st.session_state.search_type = "Fast (up to 5 sources)"
        if 'search_results' not in st.session_state:
            st.session_state.search_results = None

        # Use session state for the input fields
        query = st.text_input("Enter your search query:", key="search_query_input")
        model_choice = st.selectbox("Choose AI model (Optional):", ["gemini", "gpt4", "claude"], key="search_model_choice")
        search_type = st.radio("Search Type:", ["Fast (up to 5 sources)", "Deep (up to 10 sources)"], key="search_type")

        async def run_search():
            search_type_param = "fast" if st.session_state.search_type == "Fast (up to 5 sources)" else "deep"
            current_date = datetime.now().strftime("%Y-%m-%d")
            query_with_date = f"{st.session_state.search_query} - {current_date}"
            
            websites_used, youtube_videos_used, combined_content, full_response, word_count = await search_and_summarize(
                query_with_date, 
                st.session_state.search_model_choice, 
                search_type_param, 
                update_progress
            )
            
            st.session_state.search_results = {
                "websites_used": websites_used,
                "youtube_videos_used": youtube_videos_used,
                "combined_content": combined_content,
                "full_response": full_response,
                "word_count": word_count
            }

        if st.button("Search and Summarize"):
            if query:
                # Update session state
                st.session_state.search_query = query

                status_text = st.empty()
                progress_bar = st.progress(0)

                async def update_progress(message, progress):
                    status_text.text(message)
                    progress_bar.progress(progress)

                # Create and run the task
                search_task = asyncio.create_task(run_search())
                
                # Wait for the task to complete
                st.spinner("Searching and summarizing...")
                await search_task

                # Rerun the app to display results
                st.rerun()

        # Display results if available
        if st.session_state.search_results:
            st.write(f"Search completed. Used {st.session_state.search_results['websites_used']} websites and {st.session_state.search_results['youtube_videos_used']} YouTube videos.")
            st.write(f"Combined content word count: {st.session_state.search_results['word_count']}")

            # Show combined content in an expandable box
            with st.expander("Show Combined Content"):
                st.text_area("Combined Content", value=st.session_state.search_results['combined_content'], height=300)

            # Display the full response
            st.subheader("Search Results:")
            st.markdown(st.session_state.search_results['full_response'])

        elif query:
            st.warning("Please click 'Search and Summarize' to start the search.")
        else:
            st.warning("Please enter a search query.")

    elif tool_choice == "Document Chat Assistant":
        st.header("Document Chat Assistant 📚")
        
        st.write("""
        Welcome! This app helps you have a conversation with your documents. Here's how to use it:
        
        1. **Upload Documents**: Use the file uploader below to add your PDF files
        2. **Start Chatting**: Once processed, simply type your questions about the documents
        3. **Get Answers**: The assistant will provide detailed responses based on your documents' content
        4. **Save Your Chat**: Remember to click 'Export Chat' to save your conversation - chats are not permanently stored!
        
        Already uploaded files will appear in the dropdown menu below. Choose one to start chatting!
        """)
        # Add warning about image limitations
        st.warning("""
            ⚠️ **Warnings**: 
            - This assistant can only process text content from PDFs, not images. 
            - Conversations are temporary and will be delete when you close or refresh the page
        """)

        # Add model and parameter settings at the beginning
        model_name = 'gpt-4o'
        k = 4

        # Add verbose toggle in sidebar
        with st.sidebar:
            verbose = st.toggle("Debug Mode (Verbose)", value=False)

        # File upload section with custom name input
        uploaded_files = st.file_uploader("Upload PDF files", type="pdf", accept_multiple_files=True)

        if uploaded_files:
            custom_name = st.text_input("Give these files a name (optional)", 
                                      help="Enter a custom name for this set of files")
            
            process_button = st.button("Process Files")
            if process_button:
                with st.spinner("Processing uploaded PDFs... This may take a while."):
                    try:
                        cache_file_name = process_uploaded_files(uploaded_files, custom_name)
                        st.success("PDFs processed successfully!")
                        # Set the selected folder to the newly processed file
                        st.session_state.rag_selected_folder = cache_file_name
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error processing PDFs: {str(e)}")
                        return

        # Get available folders
        folder_options = get_cache_folders()
        
        if not folder_options:
            st.error("No cached files found in the ./cache directory.")
            return

        # Initialize session states for RAG specifically
        if 'rag_selected_folder' not in st.session_state:
            default_file = "2025 redhorse calandar_7018873614275482397.json"
            st.session_state.rag_selected_folder = default_file if default_file in folder_options else folder_options[0]
        if 'rag_conversation_history' not in st.session_state:
            st.session_state.rag_conversation_history = []

        # Add callback function for RAG folder selection
        def on_rag_folder_change():
            if st.session_state.rag_folder_selector != st.session_state.rag_selected_folder:
                st.session_state.rag_selected_folder = st.session_state.rag_folder_selector
                st.session_state.rag_conversation_history = []  # Reset only RAG conversation history

        selected_folder = st.selectbox(
            "Select file to chat with:", 
            folder_options, 
            key='rag_folder_selector',
            index=folder_options.index(st.session_state.rag_selected_folder),
            on_change=on_rag_folder_change
        )

        # Initialize the retriever with the selected cache file
        cache_file_path = os.path.join("./cache", st.session_state.rag_selected_folder)
        with st.spinner("Initializing system... This may take a while for large files, ~10s for 1000 pages."):
            retriever = initialize_retriever_from_cache(cache_file_path)

        # Process query
        query = st.chat_input("Type your message here...")

        # Display conversation history first
        for user_msg, ai_msg in st.session_state.rag_conversation_history:
            st.chat_message("user").write(user_msg)
            st.chat_message("assistant").write(ai_msg)

        if query:
            # Display user message
            with st.chat_message("user"):
                st.write(query)
            
            # Display assistant's streaming response
            with st.chat_message("assistant"):
                message_placeholder = st.empty()
                full_response = ""
                display_filename = os.path.basename(st.session_state.rag_selected_folder).replace('.json', '')
                for chunk in process_query(query, retriever, k, st.session_state.rag_conversation_history, filename=display_filename, verbose=verbose):
                    full_response += chunk
                    message_placeholder.markdown(full_response + "▌")
            
            # Add new message to history and trim if needed
            st.session_state.rag_conversation_history.append((query, full_response))
            st.session_state.rag_conversation_history = trim_conversation_history(
                st.session_state.rag_conversation_history, 
                max_words=50000
            )

        # Add word count display
        if st.session_state.rag_conversation_history:
            total_words = sum(len(msg.split()) + len(resp.split()) 
                            for msg, resp in st.session_state.rag_conversation_history)
            st.caption(f"Current conversation length: {total_words:,} words")

        # Add New Chat button
        if st.button("New Chat"):
            st.session_state.rag_conversation_history = []
            st.rerun()

        # Export chat button
        if st.session_state.rag_conversation_history:
            col1, col2 = st.columns([1, 5])
            with col1:
                export_data, timestamp = export_conversation(st.session_state.rag_conversation_history)
                st.download_button(
                    label="Export Chat",
                    data=export_data,
                    file_name=f"conversation_{timestamp}.txt",
                    mime="text/plain"
                )


if __name__ == "__main__":
    asyncio.run(streamlit_main())